import os
import difflib
from gpt4all import GPT4All
from docx import Document

# ============================
# 1️⃣ Paths
# ============================
BASE_DIR = r"C:\Users\105393\OneDrive - Dis-Chem Pharmacies\Documents\GK_Version_compare"
MODEL_PATH = os.path.join(BASE_DIR, "models", "Meta-Llama-3-8B-Instruct.Q4_1.gguf")
OLD_FOLDER = os.path.join(BASE_DIR, "data", "version_1")
NEW_FOLDER = os.path.join(BASE_DIR, "data", "version_2")
REPORTS_FOLDER = os.path.join(BASE_DIR, "results", "reports")
os.makedirs(REPORTS_FOLDER, exist_ok=True)

REPORT_FILE = os.path.join(REPORTS_FOLDER, "summary_report.docx")

# ============================
# 2️⃣ Load GPT4All Model
# ============================
print("🔄 Loading GPT4All model...")
model = GPT4All(MODEL_PATH)  # CPU-only
print("✅ Model loaded successfully!")

# ============================
# 3️⃣ Compare Files Function
# ============================
def compare_files(file1, file2):
    """Return unified diff between two text files."""
    with open(file1, 'r', encoding='utf-8') as f1, open(file2, 'r', encoding='utf-8') as f2:
        text1 = f1.readlines()
        text2 = f2.readlines()
    diff = difflib.unified_diff(text1, text2, fromfile=file1, tofile=file2)
    return ''.join(diff)

# ============================
# 4️⃣ Summarize Diff Function
# ============================
def summarize_diff(diff_text, filename):
    """Summarize what changed and why it matters using GPT4All."""
    diff_lines = [line for line in diff_text.splitlines() if line.startswith('+') or line.startswith('-')]
    diff_snippet = "\n".join(diff_lines)[:4000]  # limit for speed

    if not diff_snippet:
        return f"No significant changes detected in {filename}."

    prompt = f"""
    Summarize the following differences in GK Software file '{filename}'.
    Explain what changed and why it might matter:

    {diff_snippet}
    """

    with model.chat_session():
        response = model.generate(prompt, max_tokens=300)
    return response

# ============================
# 5️⃣ Main Execution
# ============================
old_files = {f.lower(): f for f in os.listdir(OLD_FOLDER) if f.lower().endswith(".xml")}
new_files = {f.lower(): f for f in os.listdir(NEW_FOLDER) if f.lower().endswith(".xml")}
common_files = set(old_files.keys()).intersection(set(new_files.keys()))

if not common_files:
    print("❌ No matching XML files found in both folders.")
else:
    # Create Word document
    doc = Document()
    doc.add_heading("GK Software Version Comparison Summary", level=0)

    for fname_lower in common_files:
        file_old = os.path.join(OLD_FOLDER, old_files[fname_lower])
        file_new = os.path.join(NEW_FOLDER, new_files[fname_lower])

        print(f"\n🔍 Comparing file: {old_files[fname_lower]}")
        diff_text = compare_files(file_old, file_new)

        if diff_text.strip():
            print("✅ Differences found. Generating summary...")
            summary = summarize_diff(diff_text, old_files[fname_lower])
        else:
            print("No differences found.")
            summary = f"No differences detected in {old_files[fname_lower]}."

        # Add summary to Word document
        doc.add_heading(old_files[fname_lower], level=1)
        doc.add_paragraph(summary)

    # Save Word report
    doc.save(REPORT_FILE)
    print(f"\n💾 All summaries saved to: {REPORT_FILE}")
